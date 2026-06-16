from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\andre\Desktop\EURU TOS MAIN")
OUT = ROOT / "output" / "pdf"
OUT.mkdir(parents=True, exist_ok=True)
DOCX_PATH = OUT / "Caso_Practico_Final_Ciberseguridad_Resuelto.docx"

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 99, 110)
LIGHT = "E8EEF5"
PALE = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        lead, rest = text[: len(bold_lead)], text[len(bold_lead) :]
        r = p.add_run(lead)
        r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_question(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, PALE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Pregunta: " + text)
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 14, 7),
    ("Heading 2", 13, BLUE, 10, 5),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.first_line_indent = Inches(-0.25)
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.1

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("Fundamentos y Estrategias Avanzadas de Ciberseguridad")
set_run_font(hr, size=8.5, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("Caso práctico final  |  Página ")
set_run_font(fr, size=8.5, color=GRAY)
add_page_field(fp)

# Opening block: editorial-cover-inspired, compact enough to preserve answer space.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(26)
p.paragraph_format.space_after = Pt(7)
r = p.add_run("CASO PRÁCTICO FINAL")
set_run_font(r, size=24, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Fundamentos y Estrategias Avanzadas de Ciberseguridad")
set_run_font(r, size=15, bold=True, color=BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Informe de seguridad digital para TechNova")
set_run_font(r, size=11.5, italic=True, color=GRAY)

meta = doc.add_table(rows=2, cols=2)
set_table_geometry(meta, [4680, 4680])
for row in meta.rows:
    for cell in row.cells:
        shade_cell(cell, LIGHT)
meta.cell(0, 0).text = "Alumno/a: ______________________________"
meta.cell(0, 1).text = "Fecha: __________________"
meta.cell(1, 0).text = "Curso: Ciberseguridad"
meta.cell(1, 1).text = "Empresa analizada: TechNova"
for row in meta.rows:
    for cell in row.cells:
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=9.5, color=NAVY)

add_heading(doc, "1. Propósito del informe", 1)
add_question(doc, "¿Qué querías resolver con tus recomendaciones?")
add_body(
    doc,
    "El propósito de este informe es identificar los principales riesgos de seguridad que afectan a "
    "TechNova durante su proceso de digitalización y proponer controles realistas para reducirlos. "
    "Las recomendaciones buscan proteger los datos personales y corporativos, evitar accesos no "
    "autorizados, conservar la exactitud de la información y asegurar que los servicios puedan "
    "continuar funcionando ante incidentes. También pretenden reforzar la responsabilidad de los "
    "empleados, mejorar la capacidad de respuesta de la empresa y facilitar el cumplimiento del "
    "Reglamento General de Protección de Datos (RGPD) y de la LOPDGDD."
)

add_heading(doc, "2. Amenazas detectadas", 1)
add_question(
    doc,
    "Explica qué riesgos afectan a la empresa y relaciónalos con la tríada de la ciberseguridad "
    "(Confidencialidad, Integridad y Disponibilidad)."
)

threats = doc.add_table(rows=1, cols=4)
set_table_geometry(threats, [1750, 2900, 2355, 2355])
headers = ["Amenaza", "Riesgo para TechNova", "Tríada CIA afectada", "Posible impacto"]
for i, text in enumerate(headers):
    cell = threats.cell(0, i)
    shade_cell(cell, "2E74B5")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(text)
    set_run_font(rr, size=9.5, bold=True, color=WHITE)

rows = [
    (
        "Correos de phishing",
        "Un empleado puede entregar credenciales, abrir un archivo malicioso o autorizar una acción fraudulenta.",
        "Confidencialidad e integridad; también disponibilidad si se instala ransomware.",
        "Robo de cuentas, filtración de datos, fraude, malware o interrupción de servicios.",
    ),
    (
        "Contraseñas débiles en la nube",
        "Facilitan ataques de fuerza bruta, reutilización de credenciales y accesos no autorizados.",
        "Confidencialidad e integridad.",
        "Lectura, modificación o eliminación de información y suplantación de usuarios.",
    ),
    (
        "Pérdida de un portátil con datos sensibles",
        "Una persona ajena puede acceder al equipo si el disco y la sesión no están protegidos.",
        "Confidencialidad.",
        "Brecha de datos personales, pérdida de propiedad intelectual y daño reputacional.",
    ),
    (
        "Ausencia de copia de seguridad",
        "La información del portátil puede quedar irrecuperable por pérdida, avería o ransomware.",
        "Disponibilidad e integridad.",
        "Paralización del trabajo, reconstrucción incompleta de datos y costes operativos.",
    ),
]
for values in rows:
    cells = threats.add_row().cells
    for idx, text in enumerate(values):
        cells[idx].text = text
        for p in cells[idx].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, size=9)
set_table_geometry(threats, [1750, 2900, 2355, 2355])

add_heading(doc, "3. Buenas prácticas organizativas", 1)
add_question(
    doc,
    "Propón al menos tres medidas organizativas para los empleados (por ejemplo, formación en "
    "phishing, contraseñas seguras o uso de VPN)."
)
add_bullet(
    doc,
    "Programa periódico de concienciación y simulaciones de phishing. Debe enseñar a comprobar "
    "remitentes, enlaces y archivos, y a comunicar mensajes sospechosos mediante un canal interno."
)
add_bullet(
    doc,
    "Política de contraseñas y accesos. Se deben usar claves largas y únicas, un gestor de "
    "contraseñas corporativo, autenticación multifactor y revisión periódica de permisos."
)
add_bullet(
    doc,
    "Procedimiento para trabajo remoto y movilidad. El acceso desde redes externas debe realizarse "
    "con VPN, evitando redes Wi-Fi públicas sin protección y bloqueando el equipo cuando no se use."
)
add_bullet(
    doc,
    "Protocolo de notificación de incidentes. Todo extravío, mensaje fraudulento o acceso extraño "
    "debe comunicarse inmediatamente al responsable de seguridad, sin ocultar errores."
)
add_bullet(
    doc,
    "Gestión del ciclo de vida de la información. Clasificar los datos, limitar su almacenamiento "
    "local y retirar los permisos y dispositivos cuando una persona cambia de función o deja la empresa."
)

add_heading(doc, "4. Aspectos legales", 1)
add_question(
    doc,
    "Señala al menos una obligación normativa que la empresa debe cumplir (por ejemplo, RGPD) y "
    "un riesgo en caso de incumplimiento."
)
add_body(
    doc,
    "TechNova debe cumplir el RGPD y, en España, la Ley Orgánica 3/2018 de Protección de Datos "
    "Personales y garantía de los derechos digitales (LOPDGDD). Una obligación esencial es aplicar "
    "medidas técnicas y organizativas apropiadas al riesgo, como exige el artículo 32 del RGPD. "
    "Además, si la pérdida del portátil supone una brecha de datos personales con riesgo para las "
    "personas, la empresa debe documentarla y notificarla a la autoridad de control sin dilación "
    "indebida y, cuando sea posible, dentro de las 72 horas desde que tenga constancia, conforme al "
    "artículo 33. Si existe un alto riesgo, también puede ser necesario informar a los afectados."
)
add_body(
    doc,
    "El incumplimiento puede provocar investigaciones, órdenes correctivas, reclamaciones, daño "
    "reputacional y sanciones administrativas. Para las infracciones más graves, el RGPD contempla "
    "multas de hasta 20 millones de euros o el 4 % del volumen de negocio total anual mundial del "
    "ejercicio anterior, optándose por la cuantía mayor. La sanción concreta dependerá de factores "
    "como la gravedad, duración, negligencia, datos afectados y cooperación con la autoridad."
)

add_heading(doc, "5. Medidas técnicas de protección", 1)
add_question(
    doc,
    "Recomienda dos soluciones de seguridad técnica (por ejemplo, firewall, antivirus, IDS o "
    "copias de seguridad) y justifica su utilidad."
)

solutions = [
    (
        "Copias de seguridad cifradas y verificadas",
        "Aplicar la regla 3-2-1: tres copias, en dos soportes y una fuera del entorno principal. "
        "Deben automatizarse, limitarse por permisos y probarse mediante restauraciones periódicas. "
        "Esto permite recuperar información tras pérdida, avería, borrado accidental o ransomware.",
    ),
    (
        "Cifrado de disco y gestión remota del portátil",
        "El cifrado completo impide leer los datos sin la clave, incluso si se extrae el disco. "
        "La gestión centralizada debe permitir bloquear o borrar el equipo, aplicar parches y "
        "mantener un inventario actualizado de dispositivos.",
    ),
    (
        "MFA, protección del correo y seguridad de endpoints",
        "La autenticación multifactor reduce el uso fraudulento de contraseñas robadas. El filtrado "
        "antiphishing y una solución antivirus/EDR ayudan a bloquear enlaces, adjuntos y procesos "
        "maliciosos, además de generar alertas para una respuesta rápida.",
    ),
    (
        "Firewall y monitorización",
        "Un firewall limita conexiones no autorizadas y segmenta recursos críticos. Los registros "
        "centralizados y un IDS permiten detectar patrones anómalos, investigar incidentes y "
        "conservar evidencias para mejorar los controles.",
    ),
]
for title, detail in solutions:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title + ". ")
    set_run_font(r, bold=True, color=NAVY)
    p.add_run(detail)

add_heading(doc, "6. Plan de aplicación recomendado", 1)
plan = doc.add_table(rows=1, cols=3)
set_table_geometry(plan, [1800, 4680, 2880])
for idx, text in enumerate(("Prioridad", "Acción", "Resultado esperado")):
    shade_cell(plan.cell(0, idx), "2E74B5")
    p = plan.cell(0, idx).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = p.add_run(text)
    set_run_font(rr, size=9.5, bold=True, color=WHITE)
plan_rows = [
    ("Inmediata", "Revocar sesiones del portátil perdido, cambiar credenciales y evaluar la brecha.", "Contener el incidente."),
    ("0-30 días", "Activar MFA, cifrado, copias 3-2-1, gestor de contraseñas y filtro antiphishing.", "Reducir los riesgos más urgentes."),
    ("30-60 días", "Formar al personal, aprobar políticas y probar restauraciones y respuesta a incidentes.", "Convertir los controles en hábitos."),
    ("Continua", "Revisar accesos, parches, alertas, inventario y resultados de simulaciones.", "Mejora sostenida y evidencias de cumplimiento."),
]
for values in plan_rows:
    cells = plan.add_row().cells
    for idx, text in enumerate(values):
        cells[idx].text = text
        for p in cells[idx].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                set_run_font(run, size=9.2)
set_table_geometry(plan, [1800, 4680, 2880])

add_heading(doc, "7. Reflexión final", 1)
add_question(
    doc,
    "En 100-150 palabras, explica qué aprendiste del curso y cómo aplicarías estas medidas en tu "
    "vida personal, académica o profesional."
)
reflection = (
    "Este curso me ha enseñado que la ciberseguridad no depende únicamente de herramientas técnicas, "
    "sino también de las decisiones diarias de las personas y de una adecuada organización. He "
    "aprendido a analizar los riesgos mediante la confidencialidad, la integridad y la disponibilidad, "
    "y a comprender que incidentes aparentemente simples, como reutilizar una contraseña o perder un "
    "portátil, pueden producir consecuencias graves. En mi vida personal utilizaré contraseñas únicas, "
    "autenticación multifactor, actualizaciones y copias de seguridad. En el ámbito académico y "
    "profesional comprobaré los mensajes antes de abrir enlaces, protegeré los dispositivos y "
    "comunicaré rápidamente cualquier incidente. También tendré en cuenta la privacidad y el tratamiento "
    "responsable de los datos. Mi principal conclusión es que prevenir, formar y responder con rapidez "
    "resulta más eficaz que actuar únicamente después de sufrir un ataque."
)
add_body(doc, reflection)
add_body(doc, "Extensión de la reflexión: 129 palabras.")

add_heading(doc, "8. Conclusión", 1)
add_body(
    doc,
    "TechNova presenta riesgos habituales pero relevantes: ingeniería social, credenciales débiles, "
    "pérdida de dispositivos y falta de respaldo. La combinación de formación, políticas claras, "
    "autenticación multifactor, cifrado, copias verificadas y monitorización permite reducir estos "
    "riesgos de forma proporcionada. La seguridad debe gestionarse como un proceso continuo, con "
    "responsables definidos, pruebas periódicas y aprendizaje después de cada incidente."
)

add_heading(doc, "9. Fuentes normativas consultadas", 1)
sources = [
    (
        "Reglamento (UE) 2016/679 (RGPD), especialmente artículos 32, 33, 34 y 83.",
        "https://eur-lex.europa.eu/eli/reg/2016/679/oj?locale=es",
    ),
    (
        "Ley Orgánica 3/2018, de 5 de diciembre, de Protección de Datos Personales y garantía de los derechos digitales.",
        "https://www.boe.es/buscar/act.php?id=BOE-A-2018-16673",
    ),
    (
        "Agencia Española de Protección de Datos: recursos sobre brechas de datos personales.",
        "https://www.aepd.es/guias-y-herramientas/herramientas/comunica-brecha-rgpd",
    ),
]
for label, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(label + " ")
    r = p.add_run(url)
    r.font.color.rgb = BLUE
    r.font.underline = True

core = doc.core_properties
core.title = "Caso práctico final de ciberseguridad resuelto"
core.subject = "Informe de seguridad digital para TechNova"
core.author = "Alumno/a"
core.keywords = "ciberseguridad, TechNova, RGPD, LOPDGDD, phishing, copias de seguridad"

doc.save(DOCX_PATH)
print(DOCX_PATH)
